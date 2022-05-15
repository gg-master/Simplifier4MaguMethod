import argparse
import os
import re
from docx import Document


class WordLoader:
    def __init__(self):
        self.doc = Document()
        self.doc.add_heading('Результат работы программы', 2)
        self.doc.add_paragraph('Будьте внимательны при копировании текста. '
                               'Советуем провести проверку вручную.', )

    def add_section(self, text, title_of_section=None):
        if title_of_section:
            self.doc.add_heading(title_of_section, 5)
        self.doc.add_paragraph(text)

    def save(self, name='MaguMethod.docx'):
        try:
            self.doc.save(name)
            print(f'[+] Файл успешно сохранен по адресу: '
                  f'{os.path.abspath(os.curdir)}\\{name}')
        except Exception:
            print('[X] ОШИБКА: Сохранить файл не удалось!')
            # self.save('1' + name)


class Simplifier:
    def __init__(self, string, sep, type_task='in'):
        print(f'[+] Получены данные: {string}; {sep}; {type_task};')
        print('[+] Начало обработки введенных данных...')
        self.w_loader = WordLoader()

        self.type_task = type_task
        self.string = string
        self.sep = sep  # Разделитель внутри скобок
        self.elem_sep_pat = r'\w\d’' if '’' in string else r'\w\d'
        self.elem_sep = '’' if '’' in string else ''

        self.result_str = ''

        self.w_loader.add_section(
            title_of_section='Входные данные:',
            text=f'Тип решаемой задачи: {self.type_task}\n'
                 f'Введенная строка: "{self.string}"\n'
                 f'Указанный разделитель в скобках: "{self.sep}"\n')

        self.mul_brackets()
        self.join_subsets()
        self.convert_line_into_sets()

        self.w_loader.save()

    def mul_brackets(self) -> None:
        """
        Раскрывает скобки
        :return:
        """
        res = []
        # Раскрытие скобок
        for i in re.split('[()]', self.string):
            # Пустые элементы пропускаем
            if not i.strip():
                continue
            # При первом проходе записываем в res элементы из первй скобки
            if not res:
                res.extend([''.join(j.split()) for j in i.split(self.sep)])
            # При остальный прохоодах перемножаем скобки
            else:
                curr_left = res.copy()
                curr_right = [''.join(j.split()) for j in i.split(self.sep)]
                res.clear()
                for curr_left_el in curr_left:
                    for curr_right_el in curr_right:
                        # Откидываем повторяющиеся элементы сразу
                        if curr_left_el in curr_right_el:
                            res.append(f'{curr_right_el}')
                        elif curr_right_el in curr_left_el:
                            res.append(f'{curr_left_el}')
                        else:
                            res.append(f'{curr_left_el}{curr_right_el}')

        # Первичное упрощение
        out_res = [''.join(sorted(
            list(set(list(filter(
                lambda x: x, re.findall(self.elem_sep_pat, i)))))))
            for i in res]

        self.result_str = f' {self.sep} '.join(out_res)
        self.w_loader.add_section(title_of_section='--- Первичное упрощение ('
                                  'можно не копировать) --- ',
                                  text=self.result_str)
        # Вторичное упрощение (удаление повторяющихся элементов)
        self.result_str = f' {self.sep} '.join(sorted(list(set(out_res))))
        self.w_loader.add_section(title_of_section=
                                  '--- Вторичное упрощение: раскрыты скобки,'
                                  ' повторяющиеся элементы удалены. --- ',
                                  text=self.result_str)

    def join_subsets(self) -> None:
        """
        Убирает те множества, в которые входят другие множества.
        Исп. правило алгебры логики
        :return:
        """
        arr = self.result_str.split(f' {self.sep} ')
        out_arr = []
        for i in arr:
            f = True
            for curr_el in arr:
                if i == curr_el:
                    continue
                ss = re.findall(self.elem_sep_pat, curr_el)
                if all(map(lambda x: x in i, ss)):
                    f = False
                    break
            if f:
                out_arr.append(i)

        self.result_str = f' {self.sep} '.join(out_arr)
        self.w_loader.add_section(title_of_section=
                                  '--- Удаление лишних множеств по правилам '
                                  'алгебры логики. --- ', text=self.result_str)

    def convert_line_into_sets(self) -> None:
        """
        Преобразовывает строку со всеми мнодествами в обособленные множества
        :return:
        """
        out_arr = []
        index = 1
        for el in self.result_str.split(self.sep):
            ss = re.findall(self.elem_sep_pat, el)
            out = []
            for num in range(1, 10):
                test_s = f'Y{num}{self.elem_sep}'
                if (self.type_task == 'ex' and test_s in ss) or \
                        (self.type_task == 'in' and test_s not in ss):
                    out.append(f'x{num}')
            if out:
                out_s = f'U{index} = {{{", ".join(out)}}}, |U{index}| = ' \
                        f'{len(out)}'
                out_arr.append(out_s)
                index += 1
        self.result_str = '\n'.join(out_arr)

        type_task = "внутренней" if self.type_task == "in" else "внешней" \
            if self.type_task == "ex" else "НЕВЕРНО УКАЗАН ТИП"
        self.w_loader.add_section(
            title_of_section=f'--- Полученные множества для задачи '
            f'типа "{type_task} " устойчивости ---', text=self.result_str)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Магу, а может и не Магу. Крч для 10 номера семы.')
    parser.add_argument('sep_in_brackets', type=str,
                        help='Разделитель в скобках')
    parser.add_argument('type_task', type=str, choices=['in', 'ex'],
                        help='Тип решаемой задачи: внутренняя (in) '
                             'или внешняя (ex) устойчивость')
    parser.add_argument('source_str', type=str, help='Исходная строка.')

    args = parser.parse_args()

    s = args.source_str
    sep = args.sep_in_brackets
    tt = args.type_task

    Simplifier(s, sep, tt)
